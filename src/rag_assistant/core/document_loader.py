"""Document loader supporting multiple formats: PDF, TXT, CSV, MD, DOCX."""

from __future__ import annotations

import csv
import json
import logging
from pathlib import Path

from langchain_core.documents import Document

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Load documents from various formats."""

    SUPPORTED_FORMATS = {".txt", ".pdf", ".csv", ".md", ".docx"}

    @staticmethod
    def _load_metadata_index(directory: Path) -> dict[str, dict]:
        """
        Load metadata_index.json from a directory if present.

        Returns a dict keyed by filename, e.g.:
            {"medical_0000.txt": {"topic": "what_is_hypertension", ...}}
        """
        index_path = directory / "metadata_index.json"
        if not index_path.exists():
            return {}
        try:
            with open(index_path) as f:
                entries = json.load(f)
            return {e["filename"]: e for e in entries}
        except Exception as e:
            logger.warning(f"Could not load metadata index: {e}")
            return {}

    @staticmethod
    def load_documents(directory: str | Path) -> list[Document]:
        """
        Load all supported documents from a directory.

        If a metadata_index.json sidecar exists in the directory, its entries
        are merged into each document's metadata, enabling meaningful filtering
        (e.g. by topic) even when filenames are opaque like medical_0000.txt.

        Args:
            directory: Path to directory containing documents

        Returns:
            List of langchain Document objects
        """
        directory = Path(directory)
        if not directory.exists():
            logger.warning(f"Directory does not exist: {directory}")
            return []

        metadata_index = DocumentLoader._load_metadata_index(directory)

        documents = []
        for file_path in directory.rglob("*"):
            if (
                file_path.is_file()
                and file_path.suffix.lower() in DocumentLoader.SUPPORTED_FORMATS
            ):
                try:
                    docs = DocumentLoader.load_file(file_path)
                    if metadata_index and file_path.name in metadata_index:
                        extra = {
                            k: v
                            for k, v in metadata_index[file_path.name].items()
                            if k != "filename"
                        }
                        for doc in docs:
                            doc.metadata.update(extra)
                    documents.extend(docs)
                    logger.info(f"Loaded {len(docs)} documents from {file_path}")
                except Exception as e:
                    logger.error(f"Failed to load {file_path}: {e}")

        logger.info(f"Total documents loaded: {len(documents)}")
        return documents

    @staticmethod
    def load_file(file_path: str | Path) -> list[Document]:
        """
        Load a single document file.

        Args:
            file_path: Path to file

        Returns:
            List of Document objects
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".txt":
            return DocumentLoader._load_txt(file_path)
        elif suffix == ".pdf":
            return DocumentLoader._load_pdf(file_path)
        elif suffix == ".csv":
            return DocumentLoader._load_csv(file_path)
        elif suffix == ".md":
            return DocumentLoader._load_md(file_path)
        elif suffix == ".docx":
            return DocumentLoader._load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    @staticmethod
    def _load_txt(file_path: Path) -> list[Document]:
        """Load text file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return [
            Document(
                page_content=content,
                metadata={"source": file_path.name, "format": "txt"},
            )
        ]

    @staticmethod
    def _load_pdf(file_path: Path) -> list[Document]:
        """Load PDF file."""
        if PdfReader is None:
            raise ImportError(
                "pypdf is required for PDF support. Install with: pip install pypdf"
            )

        documents = []
        reader = PdfReader(file_path)

        for page_num, page in enumerate(reader.pages):
            content = page.extract_text()
            documents.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": file_path.name,
                        "format": "pdf",
                        "page": page_num + 1,
                    },
                )
            )

        return documents

    @staticmethod
    def _load_csv(file_path: Path) -> list[Document]:
        """Load CSV file."""
        documents = []
        with open(file_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row_num, row in enumerate(reader, start=1):
                content = " ".join([f"{k}: {v}" for k, v in row.items()])
                documents.append(
                    Document(
                        page_content=content,
                        metadata={
                            "source": file_path.name,
                            "format": "csv",
                            "row": row_num,
                        },
                    )
                )

        return documents

    @staticmethod
    def _load_md(file_path: Path) -> list[Document]:
        """Load Markdown file."""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        return [
            Document(
                page_content=content,
                metadata={"source": file_path.name, "format": "md"},
            )
        ]

    @staticmethod
    def _load_docx(file_path: Path) -> list[Document]:
        """Load DOCX file."""
        if DocxDocument is None:
            raise ImportError(
                "python-docx is required for DOCX support. Install with: pip install python-docx"
            )

        doc = DocxDocument(file_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        return [
            Document(
                page_content=content,
                metadata={"source": file_path.name, "format": "docx"},
            )
        ]
